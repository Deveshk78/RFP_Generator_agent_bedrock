"""Export RFP markdown-style text to a formatted Word (.docx) document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def _add_cover_page(
    doc: Document,
    *,
    title: str,
    organization: str,
    domain_label: str | None,
    category: str | None,
) -> None:
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("REQUEST FOR PROPOSAL")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(18)

    doc.add_paragraph()
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.add_run(f"Issuing Organization: {organization}").font.size = Pt(12)

    if domain_label:
        dom_p = doc.add_paragraph()
        dom_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dom_p.add_run(f"Industry Domain: {domain_label} ({category or 'General'})").font.size = Pt(11)

    doc.add_page_break()


def _add_formatted_line(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    if stripped.startswith("#### "):
        p = doc.add_heading(stripped[5:], level=4)
        p.style.font.size = Pt(11)
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        return
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
        return

    numbered = re.match(r"^(\d+)\.\s+(.+)", stripped)
    if numbered:
        doc.add_paragraph(numbered.group(2), style="List Number")
        return

    if stripped.startswith(("- ", "* ", "• ")):
        doc.add_paragraph(stripped[2:], style="List Bullet")
        return

    if stripped.startswith("**") and stripped.endswith("**"):
        p = doc.add_paragraph()
        run = p.add_run(stripped.strip("*"))
        run.bold = True
        return

    doc.add_paragraph(stripped)


def export_rfp_to_docx(
    content: str,
    output_path: Path,
    *,
    title: str,
    organization: str,
    domain_label: str | None = None,
    category: str | None = None,
) -> Path:
    """Write a detailed Word document from RFP text content."""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_cover_page(
        doc,
        title=title,
        organization=organization,
        domain_label=domain_label,
        category=category,
    )

    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "This document contains the full Request for Proposal including scope, "
        "requirements, evaluation criteria, and submission instructions."
    )
    doc.add_page_break()

    for line in content.splitlines():
        _add_formatted_line(doc, line)

    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— End of RFP Document —")
    fr.italic = True
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(str(output_path))
    return output_path
